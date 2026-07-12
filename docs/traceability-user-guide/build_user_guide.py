from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Uninus Calendar 產銷履歷輔助系統使用操作說明.docx"
ASSETS = ROOT / "assets"

BLUE = "175CD3"; NAVY = "17324D"; LIGHT = "EAF2FF"; GREEN = "DFF3E6"; AMBER = "FFF4D6"; RED = "FCE2E0"; GRAY = "F3F5F7"

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

def set_cell_text(cell, text, bold=False, color=None, size=9):
    cell.text = ""; p = cell.paragraphs[0]; r = p.add_run(str(text)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

def add_table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = "Table Grid"
    for i,h in enumerate(headers): set_cell_text(t.rows[0].cells[i], h, True, "FFFFFF", 9); shade(t.rows[0].cells[i], NAVY)
    for ri,row in enumerate(rows):
        cells=t.add_row().cells
        for i,v in enumerate(row): set_cell_text(cells[i], v, False, None, 8.5)
        if ri%2: [shade(c, "F7F9FB") for c in cells]
    return t

def callout(doc, title, text, kind="info"):
    fills={"info":LIGHT,"success":GREEN,"warning":AMBER,"danger":RED}; t=doc.add_table(rows=1,cols=1); t.style="Table Grid"; shade(t.cell(0,0),fills[kind]);
    p=t.cell(0,0).paragraphs[0]; r=p.add_run(title+"\n"); r.bold=True; r.font.color.rgb=RGBColor.from_string(NAVY); r2=p.add_run(text); r2.font.size=Pt(9)

def step(doc, n, title, body):
    p=doc.add_paragraph(); p.style="List Number"; r=p.add_run(title+"："); r.bold=True; p.add_run(body)

def picture(doc, filename, caption):
    p=ASSETS/filename
    if p.exists():
        doc.add_picture(str(p), width=Inches(6.65)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=doc.add_paragraph(caption); c.alignment=WD_ALIGN_PARAGRAPH.CENTER; c.runs[0].italic=True; c.runs[0].font.size=Pt(8); c.runs[0].font.color.rgb=RGBColor(90,100,110)

def h(doc, text, level=1):
    doc.add_heading(text, level=level)

def bullets(doc, items):
    for x in items: doc.add_paragraph(x, style="List Bullet")

doc=Document(); sec=doc.sections[0]; sec.top_margin=Inches(.65); sec.bottom_margin=Inches(.65); sec.left_margin=Inches(.7); sec.right_margin=Inches(.7)
styles=doc.styles
styles['Normal'].font.name='Microsoft JhengHei'; styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei'); styles['Normal'].font.size=Pt(10)
for name,size,color in [('Title',26,NAVY),('Heading 1',18,NAVY),('Heading 2',14,BLUE),('Heading 3',11,NAVY)]:
    s=styles[name]; s.font.name='Microsoft JhengHei'; s._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei'); s.font.size=Pt(size); s.font.color.rgb=RGBColor.from_string(color)

# cover
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("UNINUS TRACEABILITY").bold=True; p.runs[0].font.color.rgb=RGBColor.from_string(BLUE); p.runs[0].font.size=Pt(13)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Uninus Calendar\n產銷履歷輔助系統\n使用操作說明"); r.bold=True; r.font.size=Pt(28); r.font.color.rgb=RGBColor.from_string(NAVY)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run("農場主資料 × 生產週期 × 農務作業 × 自動佐證 × AI 審閱 × 可稽核匯出").font.size=Pt(12)
doc.add_paragraph("\n")
add_table(doc,["文件資訊","內容"],[("適用版本","Uninus Calendar Service Scheduler v0.5.6"),("文件日期",str(date.today())),("適用對象","農場經營者、現場作業人員、產銷履歷管理者、內部稽核與系統管理者"),("文件性質","操作手冊與 E2E 作業規範")])
callout(doc,"重要聲明","本系統是產銷履歷資料的蒐集、關聯、完整性檢查與匯出輔助工具，不取代農業部產銷履歷（TAFT）、TGAP 驗證、第三方稽核或主管機關正式申請程序。正式送審前仍應依當期法規與驗證機構要求檢核。","warning")
doc.add_page_break()

h(doc,"文件導讀",1)
bullets(doc,["第 1–3 章：理解產銷履歷與系統資料鏈。","第 4–9 章：逐頁操作與每個欄位的用途、填寫方式與範例。","第 10 章：完整 E2E 流程，從建立農場到封存匯出。","第 11–13 章：資料治理、AI 佐證、安全注意事項與疑難排解。"])
callout(doc,"建議閱讀方式","首次導入請依第 10 章從頭操作；日常使用者可直接查閱農務作業、佐證與匯出章節。技術 ID（farm_id、cycle_id、operation_id、hash）主要供查核與系統整合，不需人工編造。")

h(doc,"1. 產銷履歷是什麼",1)
doc.add_paragraph("產銷履歷的核心不是單一表格，而是一條可追溯的證據鏈：誰、在哪裡、對哪一批產品、在什麼時間、做了什麼、使用了什麼投入品、留下哪些佐證，以及資料是否可被驗證。")
add_table(doc,["追溯問題","系統中的資料"],[
("誰生產／誰執行","農場經營者、農務作業操作人"),("在哪裡生產","農場 Farm、場區 Plot"),("哪一批產品","生產週期 Crop Cycle、批號、追溯碼"),("何時做了什麼","農務作業 AgriOperation、Calendar event"),("使用哪些資材或水源","資材/水源、數量、單位、備註"),("環境與設備當時狀況","Observed entities、start/end snapshot、state changes"),("是否有照片、文件或原始資料","EvidenceRecord、Raw Evidence Bundle、URI、content_hash"),("資料是否完整可交付","一致性檢查、integrity checks、JSON Package、CSV")])

h(doc,"2. 系統如何輔助產生產銷履歷資料",1)
doc.add_paragraph("系統把日常農務操作轉換成結構化、可關聯、可匯出的紀錄；重點是『作業當下順手記錄』，而不是事後重新拼湊。")
add_table(doc,["階段","系統輔助","產出"],[
("建立基礎資料","建立農場、場區與生產週期","farm → plot → cycle 主資料鏈"),("規劃與執行作業","農務作業同步顯示於 Calendar","AgriOperation + Calendar UID"),("宣告觀察與控制範圍","Operation Profile 集中設定 entities、Actions 與 evidence policy","可重用 Profile"),("開始作業","擷取開始狀態，必要時執行允許的開始 Action","Evidence Session / start_snapshot"),("結束作業","擷取結束狀態、服務呼叫與狀態變化並封存 hash","Raw Evidence Bundle"),("人工或 AI 整理","AI 僅建立可審閱草稿，農場人員接受或退回","ai_summary_draft + review_status"),("交付前治理","掃描缺漏、孤立佐證、失效 Calendar 關聯","integrity / consistency report"),("匯出與封存","依週期輸出 JSON Package 與 CSV","可備份、審閱、後處理的資料包")])
callout(doc,"不可混淆的治理邊界","Raw Evidence Bundle 是不可覆寫的原始來源；AI 摘要是衍生敘事，必須指向 source_session_id 與 source_raw_evidence_hash，且需人工 review。AI 不應直接控制農場設備。","danger")

h(doc,"3. 系統架構與日常入口",1)
picture(doc,"01-main-calendar-sidebar.png","圖 1　Uninus Calendar 主畫面：Calendar、產銷履歷摘要與工作台入口")
add_table(doc,["畫面元件","用途","使用時機"],[
("Calendar 清單","選擇顯示哪些 HA Calendar","確認農務作業是否出現在月/週/日視圖"),("重新整理","重新載入 Calendar 與整合狀態","新增或修改資料後同步畫面"),("產銷履歷摘要","顯示週期、作業、佐證、待辦與匯出狀態","每天快速檢查資料健康度"),("＋ 農務作業","快速新增農務作業與 Calendar 紀錄","現場完成定植、灌溉、施肥等工作後"),("開啟工作台","進入完整的履歷資料治理介面","管理主資料、佐證、Profile、一致性與匯出"),("月／週／日","以時間軸檢查事件順序","查核作業日期、漏記或重複紀錄")])

h(doc,"4. 工作台總覽與履歷範圍",1)
picture(doc,"02-workbench-overview.png","圖 2　產銷履歷工作台總覽")
add_table(doc,["欄位／指標","在履歷中的作用","操作建議／示範"],[
("目前履歷範圍","決定工作台現在檢視的農場與週期","日常處理單一批次時選特定週期；跨批盤點時選全部"),("農場／場區／週期數","確認主資料鏈是否建立","新批次開始前至少應有 1 農場、1 場區、1 週期"),("作業數","該範圍內的農務活動筆數","筆數異常為 0 時檢查篩選或 Calendar 同步"),("佐證數","可追溯證據筆數","重要作業至少保留一種佐證"),("匯出前檢查","驗證主資料、關聯、hash 與 rows 範圍","全部綠色再交付；警告需回原頁修正"),("最近作業／最近佐證","快速定位最新紀錄","作業完成後用此處確認是否已同步")])

h(doc,"5. 基礎資料：農場、場區與生產週期",1)
picture(doc,"03-workbench-master-data.png","圖 3　基礎資料管理：以階層方式維護農場、場區與週期")
h(doc,"5.1 農場 Farm 欄位",2)
add_table(doc,["欄位","作用","填寫方法","示範"],[
("農場名稱","生產者／場域主要識別","使用對外正式名稱或內部唯一名稱","靜安農場"),("經營者","責任歸屬與聯絡對象","填自然人或法人名稱","王小明"),("地址","生產所在地與稽核位置","填完整行政區與地址","臺南市安平區○○路 1 號"),("電話","查核與資料確認聯絡方式","建議含區碼或手機","0912-345-678"),("狀態","控制新資料是否可繼續使用","active=啟用；inactive=停用；archived=封存","active")])
h(doc,"5.2 場區 Plot 欄位",2)
add_table(doc,["欄位","作用","填寫方法","示範"],[
("所屬農場","建立 farm → plot 關聯","先選農場，再新增場區","靜安農場"),("場區名稱","指出實際生產地點","使用現場可辨識名稱","A 區／1 號溫室"),("產品","場區主要生產品項","填常用產品名稱","小番茄"),("TGAP 類別","對應內部分類或 TGAP 類別","依實際作物類別填寫","蔬菜類"),("面積","生產規模參考","數值與單位一起填，避免只有數字","0.1 公頃"),("位置","座標、地號或文字定位","可填 GPS、地號、設施說明","北側溫室；23.00,120.20"),("狀態","場區生命週期","停耕可 inactive，歷史場區用 archived","active")])
h(doc,"5.3 生產週期 Crop Cycle 欄位",2)
add_table(doc,["欄位","作用","填寫方法","示範"],[
("場區","批次的生產地點","選正確 Plot，不要只依產品名稱判斷","A 區"),("產品","本批次產品","使用一致命名","小番茄"),("品種","區分同產品不同品種","填種苗或契作使用名稱","玉女番茄"),("批號","內部批次唯一識別","建議含日期、區域與序號","LOT-202607-A01"),("追溯碼","外部標籤或跨系統識別","需唯一；不要重複使用舊批次碼","TRC-202607-A01"),("開始日期","生產週期起點","播種、定植或批次啟動日","2026-07-01"),("預計採收日","排程與預警參考","依作物週期估算，之後可更新","2026-09-15"),("實際採收日","完成週期的實績日期","採收完成後填寫","2026-09-18"),("狀態","週期生命週期","active／inactive／archived","active")])
callout(doc,"避免重複資料","系統會檢查重複追溯碼、同場區重複批號，以及相同場區／產品／品種／開始日的疑似重複週期。遇到提示時優先選既有週期，不要改一個字重建。","warning")

h(doc,"6. Operation Profiles：觀察、控制與佐證策略",1)
doc.add_paragraph("Operation Profile 是場區層級的可重用設定，用來宣告某類農務作業要觀察哪些 Home Assistant entity、允許控制哪些 entity、開始與結束時採用哪些 Service Action，以及 Evidence Session 的取樣上限。舊版 Sensor Profile 會向下相容載入。")
add_table(doc,["欄位","作用","填寫方法","示範"],[
("Profile 名稱","辨識可重用作業範本","以場區＋用途命名","North Orchard Irrigation Sensors"),("場區","限制 Profile 適用範圍","選擇實際 Plot","北區果園"),("Observed entities","只讀觀察與佐證來源","加入溫度、濕度、流量、土壤水分等 sensor","sensor.anping_temperature"),("Control entities","允許 Action 指向的控制範圍","只加入本作業確實需要控制的 switch/valve/script","switch.irrigation_zone_1"),("預設開始 Service Action","作業開始時的允許動作","建議呼叫 script/scene，不要放複雜邏輯","script.turn_on → script.start_irrigation"),("預設結束 Service Action","作業結束時收尾或停止動作","使用明確停止 script/scene","script.turn_on → script.stop_irrigation"),("取樣間隔（秒）","Evidence Session 連續取樣頻率","不可過短；依設備變化速度設定","60"),("最大樣本數","限制單次 Session 資料量","取樣間隔×樣本數應涵蓋作業時間","120"),("最長 Session（秒）","防止 Session 無限持續","應高於正常作業時間並留緩衝","14400（4 小時）")])
callout(doc,"目前 v0.5.6 UI 狀態","Operation Profile 管理頁已可建立／編輯 Profile、entities、Actions 與 evidence policy；後端 AgriOperation 已具備 profile_id、start_actions、end_actions 欄位。但目前農務作業編輯頁尚未呈現 Operation Profile 綁定欄位。建立正式 SOP 時請將此項列為上線前檢查，避免把後端能力誤認為現場 UI 已完整開放。","warning")

h(doc,"7. 農務作業紀錄",1)
doc.add_paragraph("農務作業是履歷的主時間軸。每一筆紀錄應回答：對哪個週期、何時、由誰、做什麼、使用什麼、多少、結果如何，以及是否連結 Calendar 與佐證。")
add_table(doc,["欄位","作用","填寫方式","填寫示範"],[
("Calendar","讓作業出現在時間軸並保存 Calendar linkage","選產銷履歷專用 Calendar","calendar.chan_xiao_lu_li"),("生產週期","把作業連到場區、產品、批號","選 LOT-202607-A01，不要只看產品名","小番茄 LOT-202607-A01"),("作業類型","農務活動分類","使用一致詞彙：播種/定植、灌溉、施肥、防治、採收等","灌溉"),("排程／實際時間","建立事件與實際發生時序","現場完成後填 actual_start；計畫作業可先填 scheduled_start","2026-07-12 14:30"),("操作人","責任與追蹤","填實際執行或確認人員","王小明"),("資材／水源","投入品或來源","需可辨識，不使用『同上』","地下水／A 配方肥"),("數量","投入量或執行量","只填數值，單位另填","30"),("單位","解釋數量尺度","公升、公斤、株、分鐘等","分鐘"),("狀態","作業生命週期","planned=計畫；completed=完成；skipped=封存/略過","completed"),("備註","補充方法、異常、天候或決策依據","簡短但足以重建現場情境","滴灌 A 區；流量正常；無漏水"),("Calendar event UID","Calendar 關聯技術識別","系統產生，不人工編輯","UUID"),("record_hash","作業內容完整性識別","系統產生；修改後應更新","SHA-256 類型雜湊")])
callout(doc,"量化原則","『灌溉 30』沒有意義；應填『數量 30、單位 分鐘』或『數量 500、單位 公升』。資材欄應填水源或肥料名稱，備註補充設備區域與異常。","info")

h(doc,"8. Evidence Session 與佐證資料",1)
picture(doc,"04-workbench-evidence.png","圖 4　佐證中心：列表／圖庫、搜尋、作業綁定與詳細資料")
h(doc,"8.1 Evidence Session 自動證據鏈",2)
add_table(doc,["欄位／階段","作用","解讀方式"],[
("session_id","一次作業執行的唯一擷取視窗","所有 start/end snapshots 與 raw bundle 應指向同一 session"),("operation_id","把 Session 連到農務作業","不可空白"),("profile_id","記錄使用哪個 Operation Profile","供重建觀察與控制範圍"),("status","capturing → ready_for_ai","ready_for_ai 表示原始包已封存，可產生 AI 草稿"),("started_at / ended_at","證據視窗時間邊界","應涵蓋實際作業時間"),("start_snapshot","作業開始前 entity 狀態","可用來比較前後變化"),("end_snapshot","作業結束後 entity 狀態","與 start_snapshot 對照"),("service_calls","Session 內實際服務呼叫","查核是否執行預期 Action"),("state_changes","觀察到的狀態變化","確認設備或環境反應"),("quality","pending／complete 等品質狀態","不完整時不可當成完整佐證"),("raw_evidence_hash","原始證據包完整性雜湊","AI 摘要必須引用同一 hash")])
h(doc,"8.2 EvidenceRecord 欄位",2)
add_table(doc,["欄位","作用","填寫方法／範例"],[
("綁定農務作業","避免佐證成為孤立資料","選擇對應 operation_id"),("佐證類型","決定資料語意與顯示方式","sensor_snapshot、raw_evidence_bundle、ai_summary_draft、photo、document、note、external_uri"),("標題","讓人能快速辨識","A 區灌溉開始前流量快照"),("content","小型 JSON 或結構化內容","{\"flow\":12.4,\"unit\":\"L/min\"}"),("來源 entity","標示 HA sensor/camera 或 Session","sensor.flow_meter_a；evidence_session:<id>"),("URI／檔案參照","大型照片或文件位置","/media/trace/2026-07-12-a.jpg"),("content_hash","檢查內容是否變動","系統產生，不手動覆寫"),("created_at","佐證建立時間","系統產生")])
callout(doc,"Raw Evidence Bundle 保護","Raw Evidence Bundle 在 UI 中為唯讀，不能覆寫或刪除。若來源資料錯誤，應新增更正紀錄或重新執行一個新的 Evidence Session，而不是改寫原始證據。","danger")

h(doc,"9. AI 佐證摘要與人工審閱",1)
doc.add_paragraph("AI 可以把已封存的 Evidence Session 轉成可讀敘事，但不能取代原始資料，也不能直接成為正式佐證。系統會把 AI 輸出標記為 ai_summary_draft，並要求農場人員接受或退回。")
add_table(doc,["欄位","作用","檢核重點"],[
("narrative","AI 產生的作業摘要","不得加入原始資料不存在的推論"),("model_identity","產生草稿的模型／代理識別","需可追溯，例如 Hermes E2E validator"),("policy_version","使用的審閱政策版本","政策更新後仍能知道當時規則"),("source_session_id","引用的 Evidence Session","必須存在且已 ready_for_ai"),("source_raw_evidence_hash","引用的不可變原始 hash","必須與 Session 相同"),("review_status","pending_farmer_review／accepted／rejected","只有 accepted 才視為已審閱正式佐證"),("review_note","接受或退回原因","退回時應說明錯誤或缺漏")])
step(doc,1,"開啟草稿","在佐證中心選取 ai_summary_draft。")
step(doc,2,"比對來源","確認來源 Session、Raw Evidence Hash、時間、observed entities 與作業一致。")
step(doc,3,"閱讀敘事","檢查數值、因果、異常與結論是否忠於原始資料。")
step(doc,4,"接受或退回","正確則『接受為正式佐證』；不正確則『退回草稿』並留下原因。")

h(doc,"10. 完整 E2E 流程與填寫示範",1)
callout(doc,"情境","示範：靜安農場 A 區種植玉女小番茄，批號 LOT-202607-A01；2026-07-12 執行 30 分鐘滴灌，擷取溫度、濕度與流量，完成後產生原始佐證、AI 摘要並匯出。","success")
h(doc,"10.1 第一次建置",2)
for n,title,body in [
(1,"建立農場","農場名稱＝靜安農場；經營者＝王小明；地址＝臺南市安平區○○路 1 號；電話＝0912-345-678；狀態＝啟用。"),
(2,"建立場區","所屬農場＝靜安農場；場區＝A 區；產品＝小番茄；TGAP 類別＝蔬菜類；面積＝0.1 公頃；位置＝北側溫室。"),
(3,"建立週期","產品＝小番茄；品種＝玉女番茄；批號＝LOT-202607-A01；追溯碼＝TRC-202607-A01；開始日＝2026-07-01；預計採收＝2026-09-15。"),
(4,"建立 Operation Profile","名稱＝A 區滴灌；Observed＝溫度、濕度、流量；Control＝灌溉 script/switch；開始／結束 Action 指向安全 script；取樣 60 秒、最大 120 筆、最長 14400 秒。")]: step(doc,n,title,body)
h(doc,"10.2 每次農務作業",2)
for n,title,body in [
(1,"建立計畫或現場紀錄","選 Calendar 與 LOT-202607-A01；作業類型＝灌溉；時間＝2026-07-12 14:30；操作人＝王小明。"),
(2,"填投入與量化資料","資材／水源＝地下水；數量＝30；單位＝分鐘；備註＝A 區滴灌，開始前檢查無漏水。"),
(3,"啟動證據視窗","系統建立 Evidence Session，擷取 start_snapshot；只有 allowlist 允許的 Service Action 可執行。"),
(4,"作業進行中","依取樣間隔收集 observed entities；記錄 service_calls 與 state_changes。"),
(5,"作業結束","擷取 end_snapshot，封存 Session，計算 raw_evidence_hash，建立唯讀 Raw Evidence Bundle。"),
(6,"人工補充佐證","如有現場照片或文件，在佐證中心選對應作業、類型 photo/document，填標題與 URI。"),
(7,"AI 整理與 review","產生 ai_summary_draft；人工比對原始 hash 後接受或退回。"),
(8,"重新開啟驗證","關閉工作台再重開，確認作業、佐證、hash 與 review 狀態仍存在。")]: step(doc,n,title,body)
h(doc,"10.3 交付與封存",2)
for n,title,body in [
(1,"選擇匯出範圍","在『匯出與封存』選 LOT-202607-A01，避免把其他批次混入。"),
(2,"執行完整性檢查","確認農場、場區、週期、作業、佐證、Calendar hash、rows 範圍全部通過。"),
(3,"預覽 JSON","抽查 farm_name、plot_name、product、lot_number、operation_type、actual_start、material、quantity、unit、hash。"),
(4,"下載 JSON Package","作為完整備份、稽核與技術交換格式。"),
(5,"下載 CSV","供 Excel/試算表整理；注意 CSV 是扁平 rows，不包含完整佐證內容。"),
(6,"封存生命週期","週期結束或資料不再新增時，依內部 SOP 封存；不硬刪歷史作業。")]: step(doc,n,title,body)

h(doc,"11. 一致性檢查與安全資料治理",1)
add_table(doc,["檢查項目","代表問題","建議處理"],[
("Calendar 關聯失效","operation 指向已不存在的 Calendar event","確認事件確實刪除後，使用清除 missing linkage 修復"),("Calendar event 找不到 operation","Calendar 有 AGRI event，但 stored operation 不存在","查核是否匯入不完整或資料被清除"),("duplicate operation ID","多個事件指向同一 operation","保留正確事件，修正重複關聯"),("stale cycle","作業指向不存在的週期","不得直接匯出；先還原週期或修正作業"),("orphan 佐證","EvidenceRecord 指向不存在的 operation","查核後才使用刪除 orphan 佐證"),("hash 驗證失敗","內容與保存 hash 不一致","停止交付，保留現場，進行資料調查")])
callout(doc,"封存優先","農務作業不再使用時優先封存為 skipped；安全刪除只適用於沒有任何引用的主資料。重置全部履歷資料只能從整合 Options Flow 進行，並應有明確二次確認。","warning")

h(doc,"12. 匯出格式說明",1)
picture(doc,"05-workbench-export.png","圖 5　匯出與封存：範圍 → 完整性檢查 → Package 交付")
picture(doc,"06-workbench-export-preview-json.png","圖 6　匯出與封存：指定範圍與逐項完整性檢查")
add_table(doc,["區塊／格式","內容","適用情境"],[
("rows","農場、場區、產品、類別、批號、作業、時間、投入、狀態、record_hash","CSV 與試算表主要來源"),("evidence","佐證 metadata、content、source、URI、content_hash","稽核、照片／文件參照、感測器佐證"),("evidence_sessions","開始／結束快照、service_calls、state_changes、raw hash","重建自動證據鏈"),("summary","各類筆數與缺漏統計","快速判斷交付範圍"),("integrity","ok、warning_count、逐項 checks","交付前品質門檻"),("JSON Package","完整結構與治理資訊","正式備份、系統交換、稽核"),("CSV","扁平作業 rows","Excel、列印、簡易統計；不等同完整履歷包")])

h(doc,"13. 安全、權限與操作邊界",1)
bullets(doc,[
"Service Action 必須受 Home Assistant allowlist 控制；優先允許 script.turn_on、scene.turn_on 與明確設備服務。",
"不要允許 homeassistant.*、hassio.*、shell_command.* 或廣泛萬用規則。",
"Control entities 只加入該場區／作業必要設備；Observed entities 與 Control entities 分離。",
"AI 只讀取封存後的證據並產生草稿，不應有設備控制權。",
"照片與大型文件使用 URI／媒體參照，不把二進位資料塞入 Calendar description。",
"技術 ID 與 hash 不由使用者手動改寫；畫面以業務名稱為主，技術資訊按需展開。",
"正式 E2E 應以真實 UI 建立、儲存、重開並核對，不只呼叫 API。"
])

h(doc,"14. 常見問題",1)
add_table(doc,["問題","可能原因","處理方式"],[
("Calendar 看不到新作業","未勾選目標 Calendar 或尚未重新整理","勾選 calendar.chan_xiao_lu_li，按重新整理"),("作業或佐證數量不符","目前週期／日期／狀態篩選限制","切換全部週期、最近 30 天或清除搜尋"),("CSV 沒有佐證 JSON","CSV 只輸出扁平 rows","改用 JSON Package 或 JSON 預覽"),("匯出前出現警告","缺主資料、作業、佐證、關聯或 hash 問題","依逐項 check 回到對應頁補齊"),("AI 草稿不能產生","Evidence Session 尚未 ready_for_ai 或沒有 raw hash","先完成 Session 並封存 Raw Evidence Bundle"),("Raw Evidence Bundle 無法刪除","系統刻意設為不可變原始證據","新增更正佐證或建立新 Session"),("不能刪除農場／場區／週期","仍有下游引用","先封存或處理週期、作業與佐證；不要硬刪"),("Operation Profile 看得到但作業頁無法綁定","v0.5.6 live UI 尚未呈現 profile binding 欄位","依發行版本確認；此功能完成前以後端資料與 Profile 管理頁為準")])

h(doc,"15. E2E 驗收清單",1)
checks=[
"主畫面可看到產銷履歷 Calendar 與摘要。","建立／選取農場後能建立場區。","建立／選取場區後能建立生產週期，批號與追溯碼唯一。","農務作業儲存後出現在作業列表與 Calendar。","關閉工作台再開啟，作業仍存在且欄位一致。","Operation Profile 可載入 observed/control entities、Actions 與 evidence policy。","Evidence Session 具有 start/end snapshot、時間與 operation/profile 關聯。","Raw Evidence Bundle 唯讀且具 raw_evidence_hash。","AI summary draft 顯示模型、policy、source session、source raw hash 與 review_status。","接受或退回 AI 草稿後，重開仍保留 review 結果。","一致性檢查無嚴重問題；修復按鈕只處理明確可修復項。","指定單一週期匯出時，rows 全部屬於該週期。","JSON Package、CSV 均可下載，且 JSON 預覽欄位正確。","整個 E2E 未使用未授權設備控制；破壞性操作有明確確認。"]
for c in checks: doc.add_paragraph("☐ "+c)

h(doc,"附錄 A：狀態與類型速查",1)
add_table(doc,["類別","值","意義"],[
("主資料狀態","active / inactive / archived","啟用／停用／封存"),("作業狀態","planned / completed / skipped","計畫／完成／封存或略過"),("Session 狀態","capturing / ready_for_ai","擷取中／原始證據已封存可供 AI"),("AI review","pending_farmer_review / accepted / rejected","待人工審閱／接受／退回"),("常用佐證類型","sensor_snapshot / raw_evidence_bundle / ai_summary_draft / photo / document / note / external_uri","感測快照／原始包／AI 草稿／照片／文件／備註／外部連結")])

h(doc,"附錄 B：文件維護資訊",1)
doc.add_paragraph("本版依專案 v0.5.6、manifest、agri.py 資料模型、v0.5.6 live Home Assistant UI 與 E2E 驗證結果重寫。畫面或欄位若在後續版本變更，應同步更新本文件的適用版本、欄位表、E2E 清單與截圖。")

# header/footer
for section in doc.sections:
    hp=section.header.paragraphs[0]; hp.text="Uninus Calendar 產銷履歷輔助系統｜v0.5.6"; hp.alignment=WD_ALIGN_PARAGRAPH.RIGHT; hp.runs[0].font.size=Pt(8); hp.runs[0].font.color.rgb=RGBColor(100,110,120)
    fp=section.footer.paragraphs[0]; fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    fp.add_run("Uninus Calendar 產銷履歷輔助系統使用操作說明　｜　第 ")
    fld=OxmlElement('w:fldSimple'); fld.set(qn('w:instr'),'PAGE'); fp._p.append(fld); fp.add_run(" 頁")
    for r in fp.runs: r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,110,120)

doc.core_properties.title="Uninus Calendar 產銷履歷輔助系統使用操作說明"
doc.core_properties.subject="v0.5.6 產銷履歷資料、Operation Profiles、Evidence Session、AI review 與完整 E2E 操作手冊"
doc.core_properties.author="Uninus / Hermes"
doc.save(OUT)
print(OUT)
